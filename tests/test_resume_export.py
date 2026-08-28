"""Resume PDF/DOCX export: rendering, filename safety, and ownership."""

from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document

from backend.schemas.schemas import ResumeVersionDetail, ResumeVersionProfile
from backend.services.application_service import apply_approval
from backend.services.resume_export_service import export_filename, render_resume_docx, render_resume_pdf
from backend.services.resume_version_service import create_resume_version
from backend.schemas.schemas import ApprovalRequest
from tests.mvp_helpers import TEST_USER_ID, insert_grounded_package, seed_materials_prerequisites


def _detail(**overrides) -> ResumeVersionDetail:
    defaults = dict(
        id="rv-abc123",
        job_id="greenhouse-xyz",
        job_title="Software Engineer",
        company="Acme Corp",
        version_number=1,
        created_at=datetime.now(timezone.utc),
        bullet_count=1,
        matches_current_profile=True,
        tailored_bullets=["Built a thing"],
        source_traceability_notes=["note"],
        profile=ResumeVersionProfile(name="Jordan Quill", email="jordan@example.com"),
    )
    defaults.update(overrides)
    return ResumeVersionDetail(**defaults)


def test_render_resume_pdf_produces_a_valid_pdf() -> None:
    pdf_bytes = render_resume_pdf(_detail())
    assert pdf_bytes[:4] == b"%PDF"


def test_render_resume_docx_contains_the_tailored_content() -> None:
    docx_bytes = render_resume_docx(
        _detail(
            tailored_bullets=["Shipped feature X"],
            profile=ResumeVersionProfile(name="Jordan Quill", skills=["Python"]),
        )
    )
    document = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Jordan Quill" in text
    assert "Shipped feature X" in text
    assert "Python" in text


def test_render_handles_xml_special_characters_without_raising() -> None:
    """reportlab's Paragraph parses its text as a small XML dialect — a
    literal '<', '>', or '&' in a scraped job title or a tailored bullet
    must render as text, not be parsed as markup or crash the renderer."""
    detail = _detail(
        job_title="Engineer <Intern> & Co-op",
        company="A & B Corp",
        tailored_bullets=["Handled >10k req/s using C++ & Go"],
    )
    pdf_bytes = render_resume_pdf(detail)
    docx_bytes = render_resume_docx(detail)
    assert pdf_bytes[:4] == b"%PDF"
    document = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Engineer <Intern> & Co-op" in text


def test_export_filename_strips_header_injection_attempt() -> None:
    """No CR/LF anywhere in the filename is the actual safety property — an
    embedded header-looking string like "Set-Cookie:" is inert text once it
    can't start a new header line, so it isn't itself something to strip."""
    detail = _detail(job_title="Engineer\r\nSet-Cookie: evil=1", company="Acme")
    name = export_filename(detail, "pdf")
    assert "\r" not in name
    assert "\n" not in name


def test_export_filename_strips_path_traversal_attempt() -> None:
    detail = _detail(job_title="../../../etc/passwd", company="Acme")
    name = export_filename(detail, "docx")
    assert "/" not in name
    assert ".." not in name


def test_export_filename_falls_back_when_stripped_to_empty() -> None:
    detail = _detail(job_title="!!!", company="???")
    name = export_filename(detail, "pdf")
    assert name.endswith(".pdf")
    assert len(name) > len(".pdf")


def test_export_filename_never_contains_a_double_quote() -> None:
    """The route embeds this filename directly inside a quoted
    Content-Disposition header value — a literal '"' here would let scraped
    job content break out of the quoted string."""
    detail = _detail(job_title='Engineer" evil="x', company="Acme")
    name = export_filename(detail, "pdf")
    assert '"' not in name


def _approved_version(session, *, public_id: str = "manual-abc123", user_id: int = TEST_USER_ID) -> str:
    job, candidate = seed_materials_prerequisites(session, public_id=public_id, user_id=user_id)
    insert_grounded_package(session, job, candidate=candidate, user_id=user_id)
    apply_approval(
        session,
        public_id,
        user_id,
        ApprovalRequest(decision="approved", eligibility_confirmed=True),
    )
    version = create_resume_version(session, job.public_id, user_id)
    return version.id


def test_export_route_returns_pdf_with_correct_headers(isolated_client) -> None:
    client, SessionLocal = isolated_client
    with SessionLocal() as db:
        version_id = _approved_version(db, user_id=client.test_user_id)

    response = client.get(f"/api/resume-versions/{version_id}/export", params={"format": "pdf"})

    assert response.status_code == 200, response.text
    assert response.headers["content-type"] == "application/pdf"
    assert ".pdf" in response.headers["content-disposition"]
    assert response.content[:4] == b"%PDF"


def test_export_route_returns_docx_with_correct_headers(isolated_client) -> None:
    client, SessionLocal = isolated_client
    with SessionLocal() as db:
        version_id = _approved_version(db, user_id=client.test_user_id)

    response = client.get(f"/api/resume-versions/{version_id}/export", params={"format": "docx"})

    assert response.status_code == 200, response.text
    assert response.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert ".docx" in response.headers["content-disposition"]
    # A .docx is a zip archive.
    assert response.content[:2] == b"PK"


def test_export_route_rejects_an_unknown_format(isolated_client) -> None:
    client, SessionLocal = isolated_client
    with SessionLocal() as db:
        version_id = _approved_version(db, user_id=client.test_user_id)

    response = client.get(f"/api/resume-versions/{version_id}/export", params={"format": "exe"})

    assert response.status_code == 422


def test_export_route_of_another_users_version_is_a_sanitized_404(isolated_client) -> None:
    client, SessionLocal = isolated_client
    with SessionLocal() as db:
        version_id = _approved_version(db, user_id=client.test_user_id)

    client.cookies.clear()
    signup = client.post("/api/auth/signup", json={"email": "user-b@example.com", "password": "a-real-password"})
    assert signup.status_code == 201, signup.text

    response = client.get(f"/api/resume-versions/{version_id}/export", params={"format": "pdf"})

    assert response.status_code == 404


def test_export_route_requires_authentication(isolated_client) -> None:
    client, SessionLocal = isolated_client
    with SessionLocal() as db:
        version_id = _approved_version(db, user_id=client.test_user_id)

    client.post("/api/auth/logout")
    response = client.get(f"/api/resume-versions/{version_id}/export", params={"format": "pdf"})
    assert response.status_code == 401
