"""Render an immutable ResumeVersion to PDF or DOCX.

Regenerates from the stored structured snapshot (ResumeVersionDetail) on
every call — never edits a binary file, never re-parses the original
upload. The only model-authored text in the output is tailored_bullets;
every other field is copied verbatim from the candidate's own stored
profile snapshot, so nothing here can introduce a claim that wasn't already
approved when the version was created.
"""

from __future__ import annotations

import io
from xml.sax.saxutils import escape as _xml_escape

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

from backend.schemas.schemas import ResumeVersionDetail
from backend.services.safe_filename import safe_filename_stem


def export_filename(detail: ResumeVersionDetail, extension: str) -> str:
    """A safe, human-readable download filename for one resume version."""
    parts = [detail.job_title, detail.company, f"v{detail.version_number}"]
    stem = safe_filename_stem(parts, default=f"resume-v{detail.version_number}")
    return f"{stem}.{extension}"


def _experience_line(entry: object) -> tuple[str, list[str]]:
    if not isinstance(entry, dict):
        return (str(entry), [])
    title = str(entry.get("title") or "").strip()
    company = str(entry.get("company") or "").strip()
    start = str(entry.get("start_date") or "").strip()
    end = str(entry.get("end_date") or "").strip()
    dates = " – ".join(part for part in (start, end) if part)
    header = " · ".join(part for part in (title, company) if part)
    if dates:
        header = f"{header} ({dates})" if header else dates
    highlights = [str(h) for h in (entry.get("highlights") or []) if str(h).strip()]
    return (header or "Experience", highlights)


def _education_line(entry: object) -> str:
    if not isinstance(entry, dict):
        return str(entry)
    institution = str(entry.get("institution") or "").strip()
    degree = str(entry.get("degree") or "").strip()
    field = str(entry.get("field") or "").strip()
    year = str(entry.get("graduation_year") or "").strip()
    degree_field = ", ".join(part for part in (degree, field) if part)
    pieces = [p for p in (institution, degree_field, year) if p]
    return " · ".join(pieces)


def _project_line(entry: object) -> tuple[str, str | None]:
    if not isinstance(entry, dict):
        return (str(entry), None)
    name = str(entry.get("name") or "").strip() or "Project"
    description = str(entry.get("description") or "").strip() or None
    return (name, description)


def render_resume_docx(detail: ResumeVersionDetail) -> bytes:
    profile = detail.profile
    document = Document()

    document.add_heading(profile.name or "Resume", level=0)

    contact_parts = [
        value
        for value in (profile.email, profile.phone, profile.linkedin_url, profile.github_url, profile.portfolio_url)
        if value
    ]
    if contact_parts:
        document.add_paragraph(" | ".join(contact_parts))

    if detail.tailored_bullets:
        document.add_heading(f"Highlights for {detail.job_title} at {detail.company}", level=1)
        for bullet in detail.tailored_bullets:
            document.add_paragraph(bullet, style="List Bullet")

    if profile.skills:
        document.add_heading("Skills", level=1)
        document.add_paragraph(", ".join(str(s) for s in profile.skills))

    if profile.experience:
        document.add_heading("Experience", level=1)
        for entry in profile.experience:
            header, highlights = _experience_line(entry)
            document.add_paragraph(header, style="Heading 3")
            for highlight in highlights:
                document.add_paragraph(highlight, style="List Bullet")

    if profile.education:
        document.add_heading("Education", level=1)
        for entry in profile.education:
            line = _education_line(entry)
            if line:
                document.add_paragraph(line)

    if profile.projects:
        document.add_heading("Projects", level=1)
        for entry in profile.projects:
            name, description = _project_line(entry)
            document.add_paragraph(name, style="Heading 3")
            if description:
                document.add_paragraph(description)

    if profile.certifications:
        document.add_heading("Certifications", level=1)
        document.add_paragraph(", ".join(str(c) for c in profile.certifications))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_resume_pdf(detail: ResumeVersionDetail) -> bytes:
    profile = detail.profile
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle("CPHeading", parent=styles["Heading2"], spaceBefore=12, spaceAfter=4)
    subheading_style = ParagraphStyle("CPSubheading", parent=styles["Heading4"], spaceBefore=6, spaceAfter=2)
    body_style = styles["BodyText"]

    # reportlab's Paragraph parses its text as a small XML/HTML dialect, and
    # every value below (name, bullets, job title/company, ...) is either
    # candidate-authored or copied from a scraped posting, not markup — a
    # literal "<", ">", or "&" must render as text, not be parsed as a tag.
    def p(text: object, style: ParagraphStyle = body_style) -> Paragraph:
        return Paragraph(_xml_escape(str(text)), style)

    story: list = [p(profile.name or "Resume", styles["Title"])]

    contact_parts = [
        value
        for value in (profile.email, profile.phone, profile.linkedin_url, profile.github_url, profile.portfolio_url)
        if value
    ]
    if contact_parts:
        story.append(p(" | ".join(contact_parts)))

    if detail.tailored_bullets:
        story.append(p(f"Highlights for {detail.job_title} at {detail.company}", heading_style))
        story.append(
            ListFlowable([ListItem(p(bullet)) for bullet in detail.tailored_bullets], bulletType="bullet")
        )

    if profile.skills:
        story.append(p("Skills", heading_style))
        story.append(p(", ".join(str(s) for s in profile.skills)))

    if profile.experience:
        story.append(p("Experience", heading_style))
        for entry in profile.experience:
            header, highlights = _experience_line(entry)
            story.append(p(header, subheading_style))
            if highlights:
                story.append(ListFlowable([ListItem(p(h)) for h in highlights], bulletType="bullet"))

    if profile.education:
        story.append(p("Education", heading_style))
        for entry in profile.education:
            line = _education_line(entry)
            if line:
                story.append(p(line))

    if profile.projects:
        story.append(p("Projects", heading_style))
        for entry in profile.projects:
            name, description = _project_line(entry)
            story.append(p(name, subheading_style))
            if description:
                story.append(p(description))

    if profile.certifications:
        story.append(p("Certifications", heading_style))
        story.append(p(", ".join(str(c) for c in profile.certifications)))

    story.append(Spacer(1, 0))
    doc.build(story)
    return buffer.getvalue()
